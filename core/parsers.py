"""
Resume parsing utilities.

- extract_text(file, filename): pulls raw text out of PDF / DOCX / TXT files.
- parse_resume(text): extracts structured data — contact info, links, skills,
  education signals, and an estimated years-of-experience figure — that feeds
  the hybrid screening engine.
"""
import io
import re

from pypdf import PdfReader
from docx import Document

from core.skills_data import SKILL_INDEX

MAX_RESUME_CHARS = 200_000

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(
    r"(?<!\w)(?:\+?\d{1,3}[\s.-]?)?(?:\(\d{2,4}\)[\s.-]?)?\d{3,4}[\s.-]?\d{3,4}[\s.-]?\d{3,4}(?!\w)"
)
URL_RE = re.compile(r"(?:https?://|www\.)[^\s]+", re.IGNORECASE)
LINKEDIN_RE = re.compile(r"(?:linkedin\.com/in/|linkedin\.com/in)[^\s]+", re.IGNORECASE)
GITHUB_RE = re.compile(r"(?:github\.com/)[^\s]+", re.IGNORECASE)
YEARS_RE = re.compile(r"(\d{1,2})\+?\s*(?:years|yrs|yr)[^a-z]", re.IGNORECASE)
DATE_RANGE_RE = re.compile(
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*20\d{2}|20\d{2})\s*[-–—to]+\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*20\d{2}|20\d{2}|present|current|now)",
    re.IGNORECASE,
)

DEGREE_KEYWORDS = [
    "bachelor", "b.s.", "b.sc", "b.tech", "b.e.", "ba in", "master", "m.s.", "m.sc",
    "m.tech", "mba", "ph.d", "phd", "doctorate", "associate degree", "b.com",
    "b.a.", "msc", "bsc", "mba", "mca", "bca",
]
UNIVERSITY_KEYWORDS = ["university", "college", "institute", "school of", "academy", "polytechnic"]
EDUCATION_SECTION_RE = re.compile(
    r"(?is)(?:education|academics?|qualifications?)[^\n]*\n(.+?)(?=\n\s*(?:experience|work|employment|skills|projects?|certifications?|languages?|interests?|references?|$))"
)


def extract_text(uploaded_file):
    """Extract raw text from an uploaded resume file."""
    filename = (getattr(uploaded_file, "name", "") or "").lower()
    data = uploaded_file.read()

    if filename.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(data))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception:
            return ""
    if filename.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        except Exception:
            return ""
    if filename.endswith((".txt", ".md", ".rtf")):
        return data.decode("utf-8", errors="ignore")
    # Unknown types — attempt a raw decode.
    return data.decode("utf-8", errors="ignore")


def _normalize(text):
    return re.sub(r"\s+", " ", text.lower())


def extract_skills(text):
    """Return the list of canonical skill names found in the text."""
    normalized = _normalize(text)
    found = set()
    for alias, canonical in SKILL_INDEX.items():
        if " " in alias:
            if alias in normalized:
                found.add(canonical)
        else:
            # token or token-prefix match, e.g. "python" matches "python3"
            if re.search(rf"\b{re.escape(alias)}[a-z0-9\-+]*\b", normalized):
                found.add(canonical)
    return sorted(found)


def extract_years_experience(text):
    """Best-effort years-of-experience estimate from explicit phrases and date ranges."""
    years = None
    m = YEARS_RE.search(text)
    if m:
        years = int(m.group(1))
    else:
        # Sum date-range spans (cap the unit ceiling).
        spans = DATE_RANGE_RE.findall(text)
        total = 0.0
        for start, end in spans:
            try:
                sy = int(re.search(r"20\d{2}", start).group())
            except AttributeError:
                continue
            if end.lower() in ("present", "current", "now"):
                ey = 2026
            else:
                try:
                    ey = int(re.search(r"20\d{2}", end).group())
                except AttributeError:
                    continue
            if 1990 <= sy <= 2026 and sy <= ey:
                total += max(0, ey - sy)
        if total:
            years = round(min(total, 40), 1)
    if years is None:
        return 0
    return years


def extract_contact(text):
    emails = list(dict.fromkeys(EMAIL_RE.findall(text)))
    phones = list(dict.fromkeys(p.strip("().- ") for p in PHONE_RE.findall(text)))
    links = list(dict.fromkeys(URL_RE.findall(text)))
    linkedin = list(dict.fromkeys(LINKEDIN_RE.findall(text)))
    github = list(dict.fromkeys(GITHUB_RE.findall(text)))
    return {
        "emails": emails[:3],
        "phones": phones[:3],
        "links": links[:5],
        "linkedin": linkedin[:2],
        "github": github[:2],
    }


def extract_education(text):
    """Return dict with degree + institution signals and the raw snippet."""
    lower = text.lower()
    degrees = [d for d in DEGREE_KEYWORDS if d in lower]
    institutions = [u for u in UNIVERSITY_KEYWORDS if u in lower]
    m = EDUCATION_SECTION_RE.search(text)
    snippet = re.sub(r"\s+", " ", m.group(1)).strip()[:300] if m else ""
    return {
        "has_education": bool(degrees or snippet),
        "degrees": degrees,
        "institutions": institutions,
        "snippet": snippet,
    }


def parse_resume(text):
    """Parse raw resume text into a structured dict stored on the Resume model."""
    text = (text or "")[:MAX_RESUME_CHARS]
    return {
        "contact": extract_contact(text),
        "skills": extract_skills(text),
        "years_experience": extract_years_experience(text),
        "education": extract_education(text),
        "word_count": len(text.split()),
    }
